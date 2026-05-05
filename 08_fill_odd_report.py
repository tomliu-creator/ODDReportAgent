# Databricks notebook source
# COMMAND ----------
# MAGIC %pip install -U python-docx

# COMMAND ----------
dbutils.library.restartPython()

# COMMAND ----------
# MAGIC %run ./_config

# COMMAND ----------
# MAGIC %run ./_utils

# COMMAND ----------
# MAGIC %md
# MAGIC ## 08 Fill the ODD report template with topic assessments and chapter summaries

# COMMAND ----------

import os
import shutil
import uuid
from docx import Document

dbutils.widgets.text("engagement_id", "odd_ssga_2025")
ENGAGEMENT_ID = dbutils.widgets.get("engagement_id").strip()

paths = engagement_paths(ENGAGEMENT_ID)
out_dir_local = paths["output_local"]
os.makedirs(out_dir_local, exist_ok=True)

metadata_rows = (
    spark.table(ODD_REPORT_METADATA_TABLE)
    .filter(F.col("engagement_id") == ENGAGEMENT_ID)
    .orderBy(F.col("load_ts").desc())
    .limit(1)
    .collect()
)
if not metadata_rows:
    raise ValueError("No ODD report metadata found; run 06_parse_odd_report first.")
metadata = metadata_rows[0].asDict()

topics = [
    r.asDict()
    for r in (
        spark.table(ODD_REPORT_TOPICS_TABLE)
        .filter(F.col("engagement_id") == ENGAGEMENT_ID)
        .orderBy("topic_order")
        .collect()
    )
]
assessments = {
    r["topic_id"]: r.asDict()
    for r in (
        spark.table(ODD_TOPIC_ASSESSMENTS_TABLE)
        .filter(F.col("engagement_id") == ENGAGEMENT_ID)
        .collect()
    )
}
chapter_summaries = {
    r["chapter_code"]: r.asDict()
    for r in (
        spark.table(ODD_CHAPTER_SUMMARIES_TABLE)
        .filter(F.col("engagement_id") == ENGAGEMENT_ID)
        .collect()
    )
}

missing_topics = [t["section_code"] for t in topics if t["topic_id"] not in assessments]
if missing_topics:
    raise RuntimeError(f"Missing topic assessments for: {missing_topics}")

docx_local = uc_dbfs_to_local_path(metadata["report_template_path"])
doc = Document(docx_local)


def _set_cell_text(cell, text: str):
    text = text or ""
    cell.text = ""
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.text = ""
    for idx, block in enumerate((text or "").split("\n\n")):
        target_para = p if idx == 0 else cell.add_paragraph()
        target_para.text = block


def _set_paragraph_style_if_available(paragraph, preferred_style: str, fallback_style: str | None = None):
    available = {style.name for style in doc.styles if getattr(style, "name", None)}
    for candidate in [preferred_style, fallback_style]:
        if candidate and candidate in available:
            paragraph.style = candidate
            return


def _fill_metadata_table():
    table = doc.tables[0]
    replacements = {
        "Mandate name:": metadata.get("mandate_name") or "",
        "Manager:": metadata.get("manager_name") or "",
        "Investment strategies:": metadata.get("investment_strategy") or "",
        "Report finalization date:": metadata.get("report_finalization_date") or "",
        "Portfolio/ODD Managers:": metadata.get("portfolio_odd_managers") or "",
        "Authors:": metadata.get("authors") or "",
    }
    for row in table.rows:
        label = normalize_text(row.cells[0].text)
        if label in replacements:
            row.cells[1].text = replacements[label]


def _assessment_block(assessment: dict) -> str:
    parts = [
        f"Risk rating: {assessment.get('risk_rating') or 'Medium'}",
    ]
    if normalize_text(assessment.get("risk_rationale")):
        parts.append(f"Risk rationale: {normalize_text(assessment['risk_rationale'])}")
    parts.append(normalize_text(assessment.get("assessment_text")))
    return "\n\n".join([part for part in parts if part])


def _fill_part1_table():
    table = doc.tables[int(metadata["part1_table_index"])]
    for topic in topics:
        assessment = assessments[topic["topic_id"]]
        row_idx = int(topic["answer_row_index"])
        if row_idx >= len(table.rows):
            raise RuntimeError(f"Answer row index {row_idx} out of range for topic {topic['section_code']}")
        _set_cell_text(table.rows[row_idx].cells[1], _assessment_block(assessment))


def _fill_part2_table_and_append_summaries():
    table = doc.tables[int(metadata["part2_table_index"])]
    for row in table.rows[1:]:
        label = normalize_text(row.cells[0].text)
        if label.lower().startswith("overall conclusion"):
            summary = chapter_summaries.get("OVERALL")
        else:
            chapter_code = label[:1]
            summary = chapter_summaries.get(chapter_code)
        if summary:
            row.cells[1].text = summary.get("rating") or row.cells[1].text

    doc.add_paragraph("")
    header = doc.add_paragraph("Part 2 Summaries")
    _set_paragraph_style_if_available(header, "Heading 2", "Heading 1")
    for chapter_code in ["A", "B", "C", "D", "OVERALL"]:
        summary = chapter_summaries.get(chapter_code)
        if not summary:
            continue
        title = summary["chapter_title"]
        sub = doc.add_paragraph(title)
        _set_paragraph_style_if_available(sub, "Heading 3", "Heading 2")
        body = []
        if summary.get("rating"):
            body.append(f"Rating: {summary['rating']}")
        if normalize_text(summary.get("summary_text")):
            body.append(normalize_text(summary["summary_text"]))
        doc.add_paragraph("\n\n".join(body))


_fill_metadata_table()
_fill_part1_table()
_fill_part2_table_and_append_summaries()

out_basename = metadata["report_template_name"].rsplit(".docx", 1)[0] + ".filled.docx"
out_local = os.path.join(out_dir_local, out_basename)
out_dbfs = f"{paths['output_dbfs']}/{out_basename}"
out_dbfs_local = out_dbfs.replace("dbfs:", "/dbfs")
workspace_output_dir = "/Workspace/Users/tomliushopping@gmail.com/notebooks_ddq/output_docs"
os.makedirs(workspace_output_dir, exist_ok=True)
out_workspace_local = os.path.join(workspace_output_dir, out_basename)

workspace_tmp_dir = "/Workspace/Users/tomliushopping@gmail.com/notebooks_ddq/.tmp"
os.makedirs(workspace_tmp_dir, exist_ok=True)
tmp_out_local = os.path.join(workspace_tmp_dir, f"{uuid.uuid4().hex}.docx")
doc.save(tmp_out_local)

final_output_path = out_dbfs
try:
    shutil.copyfile(tmp_out_local, out_dbfs_local)
except OSError as e:
    shutil.copyfile(tmp_out_local, out_workspace_local)
    final_output_path = out_workspace_local
    print(f"WARNING: could not write to volume output path ({out_dbfs_local}): {e}")
    print("Fell back to workspace-local output path.")

print("Saved:", final_output_path)
show_validation_snapshot(CATALOG, SCHEMA, ENGAGEMENT_ID)
