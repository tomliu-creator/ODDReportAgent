# Databricks notebook source
# COMMAND ----------
# MAGIC %pip install -U python-docx

# COMMAND ----------
dbutils.library.restartPython()

# COMMAND ----------
# MAGIC %run ./_config

# COMMAND ----------
# MAGIC %run ./_utils

# COMMAND ----------
# MAGIC %md
# MAGIC **Legacy Notebook**
# MAGIC
# MAGIC The default ODD workflow now loads runtime behavior from `ODDAgent.md` via
# MAGIC `06_load_odd_agent_spec.py`. Keep this notebook only as a reference for the older
# MAGIC template-parsing approach.
# MAGIC
# MAGIC ## 06 Parse the ODD report template into structured topic metadata

# COMMAND ----------

import os
import re
from docx import Document

dbutils.widgets.text("engagement_id", "odd_ssga_2025")
dbutils.widgets.text("workflow_profile", "odd_report_v1")
dbutils.widgets.text("report_template_name", "")

ENGAGEMENT_ID = dbutils.widgets.get("engagement_id").strip()
PROFILE = get_workflow_profile(dbutils.widgets.get("workflow_profile").strip() or "odd_report_v1")
REPORT_TEMPLATE_NAME = dbutils.widgets.get("report_template_name").strip() or None
topic_re = re.compile(PROFILE["topic_row_regex"])

template_docs = (
    spark.table(DOCUMENTS_TABLE)
    .filter((F.col("engagement_id") == ENGAGEMENT_ID) & (F.col("source_role") == "report_template") & (F.col("is_present") == True))
    .orderBy("file_name")
)

if REPORT_TEMPLATE_NAME:
    template_docs = template_docs.filter(F.col("file_name") == REPORT_TEMPLATE_NAME)

docs = [r.asDict() for r in template_docs.collect()]
if len(docs) != 1:
    raise ValueError(f"Expected exactly 1 report template for engagement {ENGAGEMENT_ID}; found {[d['file_name'] for d in docs]}")

doc_meta = docs[0]
docx_local = uc_dbfs_to_local_path(doc_meta["file_path_dbfs"])
print("Parsing report template:", docx_local)
doc = Document(docx_local)


def _next_prompt_text(paragraphs, start_idx: int) -> str:
    values = []
    for para in paragraphs[start_idx + 1:]:
        style = para.style.name if para.style else ""
        text = normalize_text(para.text)
        if style.startswith("Heading") and values:
            break
        if text:
            values.append(text)
    return "\n".join(values).strip()


prompt_text = ""
for idx, para in enumerate(doc.paragraphs):
    if normalize_text(para.text).lower() == "prompt":
        prompt_text = _next_prompt_text(doc.paragraphs, idx)
        break

if not prompt_text:
    raise RuntimeError("Could not extract prompt text from ODD report template.")

table0_rows = []
for row in doc.tables[0].rows:
    cells = [normalize_text(c.text) for c in row.cells[:2]]
    if len(cells) >= 2:
        table0_rows.append((cells[0], cells[1]))
report_metadata = extract_label_value_metadata(table0_rows)

risk_rows = []
for idx, row in enumerate(doc.tables[1].rows):
    label_cell = normalize_text(row.cells[0].text if len(row.cells) > 0 else "")
    desc_cell = normalize_text(row.cells[1].text if len(row.cells) > 1 else "")
    combined = desc_cell or label_cell
    label = normalize_text(combined.split(":", 1)[0]).title()
    definition = normalize_text(combined.split(":", 1)[1] if ":" in combined else combined)
    if label:
        risk_rows.append({
            "engagement_id": ENGAGEMENT_ID,
            "rating_label": label,
            "rating_definition": definition,
            "sort_order": idx,
        })

topic_rows = []
part1_table_index = 2
for row_idx in range(0, len(doc.tables[part1_table_index].rows), 2):
    topic_text = normalize_text(doc.tables[part1_table_index].rows[row_idx].cells[0].text)
    if not topic_text:
        continue
    match = topic_re.match(topic_text)
    if not match:
        raise RuntimeError(f"Could not parse Part 1 topic row text: {topic_text}")
    chapter_code, chapter_title, section_code, topic_title = match.groups()
    topic_rows.append({
        "engagement_id": ENGAGEMENT_ID,
        "topic_id": sha256_hex(f"{ENGAGEMENT_ID}||{section_code}||{topic_title}"),
        "topic_order": len(topic_rows) + 1,
        "table_index": part1_table_index,
        "topic_row_index": row_idx,
        "answer_row_index": row_idx + 1,
        "chapter_code": chapter_code,
        "chapter_title": chapter_title.strip(),
        "section_code": section_code,
        "topic_title": topic_title.strip(),
        "raw_topic_text": topic_text,
    })

if not topic_rows:
    raise RuntimeError("No ODD Part 1 topics were extracted from the report template.")

# COMMAND ----------

metadata_df = spark.createDataFrame([{
    "engagement_id": ENGAGEMENT_ID,
    "report_template_document_id": doc_meta["document_id"],
    "report_template_path": doc_meta["file_path_dbfs"],
    "report_template_name": doc_meta["file_name"],
    "mandate_name": report_metadata.get("mandate_name") or doc_meta.get("mandate_name"),
    "manager_name": report_metadata.get("manager_name") or doc_meta.get("manager_name"),
    "investment_strategy": report_metadata.get("investment_strategy") or doc_meta.get("investment_strategy"),
    "report_finalization_date": report_metadata.get("report_finalization_date") or doc_meta.get("report_finalization_date"),
    "portfolio_odd_managers": report_metadata.get("portfolio_odd_managers") or doc_meta.get("portfolio_odd_managers"),
    "authors": report_metadata.get("authors") or doc_meta.get("authors"),
    "prompt_text": prompt_text,
    "part1_table_index": 2,
    "part2_table_index": 3,
}], schema=T.StructType([
    T.StructField("engagement_id", T.StringType(), False),
    T.StructField("report_template_document_id", T.StringType(), True),
    T.StructField("report_template_path", T.StringType(), True),
    T.StructField("report_template_name", T.StringType(), True),
    T.StructField("mandate_name", T.StringType(), True),
    T.StructField("manager_name", T.StringType(), True),
    T.StructField("investment_strategy", T.StringType(), True),
    T.StructField("report_finalization_date", T.StringType(), True),
    T.StructField("portfolio_odd_managers", T.StringType(), True),
    T.StructField("authors", T.StringType(), True),
    T.StructField("prompt_text", T.StringType(), True),
    T.StructField("part1_table_index", T.IntegerType(), True),
    T.StructField("part2_table_index", T.IntegerType(), True),
])).withColumn("load_ts", F.current_timestamp())

# Use explicit schemas here as well so Spark Connect does not silently widen Python ints
# to BIGINT/LongType and then fail later when the Delta targets are defined as INT.
topics_df = spark.createDataFrame(topic_rows, schema=T.StructType([
    T.StructField("engagement_id", T.StringType(), False),
    T.StructField("topic_id", T.StringType(), False),
    T.StructField("topic_order", T.IntegerType(), False),
    T.StructField("table_index", T.IntegerType(), False),
    T.StructField("topic_row_index", T.IntegerType(), False),
    T.StructField("answer_row_index", T.IntegerType(), False),
    T.StructField("chapter_code", T.StringType(), True),
    T.StructField("chapter_title", T.StringType(), True),
    T.StructField("section_code", T.StringType(), True),
    T.StructField("topic_title", T.StringType(), True),
    T.StructField("raw_topic_text", T.StringType(), True),
])).withColumn("load_ts", F.current_timestamp())

risk_df = spark.createDataFrame(risk_rows, schema=T.StructType([
    T.StructField("engagement_id", T.StringType(), False),
    T.StructField("rating_label", T.StringType(), False),
    T.StructField("rating_definition", T.StringType(), True),
    T.StructField("sort_order", T.IntegerType(), False),
])).withColumn("load_ts", F.current_timestamp())

spark.sql(f"DELETE FROM {ODD_REPORT_METADATA_TABLE} WHERE engagement_id = '{ENGAGEMENT_ID}'")
spark.sql(f"DELETE FROM {ODD_REPORT_TOPICS_TABLE} WHERE engagement_id = '{ENGAGEMENT_ID}'")
spark.sql(f"DELETE FROM {ODD_RISK_DEFINITIONS_TABLE} WHERE engagement_id = '{ENGAGEMENT_ID}'")

metadata_df.write.mode("append").saveAsTable(ODD_REPORT_METADATA_TABLE)
topics_df.write.mode("append").saveAsTable(ODD_REPORT_TOPICS_TABLE)
risk_df.write.mode("append").saveAsTable(ODD_RISK_DEFINITIONS_TABLE)

spark.sql(f"""
MERGE INTO {DOCUMENTS_TABLE} t
USING (
  SELECT
    '{doc_meta["document_id"]}' AS document_id,
    '{(report_metadata.get("manager_name") or "").replace("'", "''")}' AS manager_name,
    '{(report_metadata.get("mandate_name") or "").replace("'", "''")}' AS mandate_name,
    '{(report_metadata.get("investment_strategy") or "").replace("'", "''")}' AS investment_strategy,
    '{(report_metadata.get("report_finalization_date") or "").replace("'", "''")}' AS report_finalization_date,
    '{(report_metadata.get("portfolio_odd_managers") or "").replace("'", "''")}' AS portfolio_odd_managers,
    '{(report_metadata.get("authors") or "").replace("'", "''")}' AS authors
) s
ON t.document_id = s.document_id
WHEN MATCHED THEN UPDATE SET
  t.manager_name = nullif(s.manager_name, ''),
  t.mandate_name = nullif(s.mandate_name, ''),
  t.investment_strategy = nullif(s.investment_strategy, ''),
  t.report_finalization_date = nullif(s.report_finalization_date, ''),
  t.portfolio_odd_managers = nullif(s.portfolio_odd_managers, ''),
  t.authors = nullif(s.authors, '')
""")

# COMMAND ----------

display(spark.table(ODD_REPORT_TOPICS_TABLE).filter(F.col("engagement_id") == ENGAGEMENT_ID).orderBy("topic_order"))
display(spark.table(ODD_RISK_DEFINITIONS_TABLE).filter(F.col("engagement_id") == ENGAGEMENT_ID).orderBy("sort_order"))

show_validation_snapshot(CATALOG, SCHEMA, ENGAGEMENT_ID)
