# Databricks notebook source
# COMMAND ----------
# MAGIC %pip install -U PyMuPDF python-docx

# COMMAND ----------
dbutils.library.restartPython()

# COMMAND ----------
# MAGIC %run ./_config

# COMMAND ----------
# MAGIC %run ./_utils

# COMMAND ----------
# MAGIC %md
# MAGIC ## 03 Parse source files (PDF + DOCX) -> `document_pages`
# MAGIC
# MAGIC PDFs are parsed page by page. The completed manager DDQ is parsed paragraph by paragraph with
# MAGIC section metadata so later retrieval can target `A0100`, `B0500`, and similar sections.

# COMMAND ----------

import os
import re
from docx import Document

try:
    import fitz  # PyMuPDF
except Exception:
    import pymupdf as fitz

dbutils.widgets.text("engagement_id", "odd_ssga_2025")
dbutils.widgets.text("workflow_profile", "odd_report_v1")
ENGAGEMENT_ID = dbutils.widgets.get("engagement_id").strip()
PROFILE = get_workflow_profile(dbutils.widgets.get("workflow_profile").strip() or "odd_report_v1")

section_re = re.compile(PROFILE["section_heading_regex"])
question_re = re.compile(PROFILE["question_numbering_regex"])

pending = (
    spark.table(DOCUMENTS_TABLE)
    .filter((F.col("engagement_id") == ENGAGEMENT_ID) & (F.col("is_present") == True))
    .filter(F.col("parse_status").isin(["pending", "error"]))
    .select("document_id", "file_name", "file_ext", "file_path_dbfs", "file_path_local", "source_role", "source_tier")
    .orderBy("source_role", "file_name")
)

pending_count = pending.count()
print("Pending documents to parse:", pending_count)
if pending_count == 0:
    show_validation_snapshot(CATALOG, SCHEMA, ENGAGEMENT_ID)
    dbutils.notebook.exit("Nothing to parse.")

docs = [r.asDict() for r in pending.collect()]


def _ddq_table_metadata(doc: Document) -> dict:
    if not doc.tables:
        return {}
    rows = []
    for row in doc.tables[0].rows:
        cells = [normalize_text(c.text) for c in row.cells[:2]]
        if len(cells) >= 2 and (cells[0] or cells[1]):
            rows.append((cells[0], cells[1]))
    return extract_label_value_metadata(rows)


def _make_page_row(
    doc_meta: dict,
    page_num: int,
    page_text: str,
    parse_method: str,
    content_role: str,
    source_page_num: int | None,
    source_para_num: int | None,
    source_locator_type: str,
    section_code: str | None = None,
    section_title: str | None = None,
    chapter_code: str | None = None,
    chapter_title: str | None = None,
    question_number: int | None = None,
    has_substantive_answer: bool | None = None,
    referenced_appendices: str | None = None,
) -> dict:
    return {
        "engagement_id": ENGAGEMENT_ID,
        "document_id": doc_meta["document_id"],
        "file_name": doc_meta["file_name"],
        "source_path_dbfs": doc_meta["file_path_dbfs"],
        "source_path_local": doc_meta["file_path_local"],
        "page_num": page_num,
        "source_page_num": source_page_num,
        "source_para_num": source_para_num,
        "source_locator_type": source_locator_type,
        "source_locator_label": build_source_locator_label(
            source_page_num=source_page_num,
            source_para_num=source_para_num,
            source_locator_type=source_locator_type,
        ),
        "page_text": page_text,
        "page_char_count": len(page_text or ""),
        "parse_method": parse_method,
        "source_role": doc_meta["source_role"],
        "source_tier": doc_meta["source_tier"],
        "section_code": section_code,
        "section_title": section_title,
        "chapter_code": chapter_code,
        "chapter_title": chapter_title,
        "content_role": content_role,
        "question_number": question_number,
        "has_substantive_answer": has_substantive_answer,
        "referenced_appendices": referenced_appendices,
    }


def _parse_manager_ddq(local_path: str, doc_meta: dict) -> tuple[list[dict], dict]:
    doc = Document(local_path)
    metadata = _ddq_table_metadata(doc)
    raw_rows = []
    current_section_code = None
    current_section_title = None
    current_chapter_code = None
    current_chapter_title = None
    current_question_number = None
    chapter_titles = {}

    for idx, para in enumerate(doc.paragraphs):
        text = normalize_text(para.text)
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        if style_name == "Heading 1":
            m_ch = re.match(r"^([A-D])\.\s+(.+)$", text)
            if m_ch:
                current_chapter_code = m_ch.group(1)
                current_chapter_title = m_ch.group(2).strip()
                chapter_titles[current_chapter_code] = current_chapter_title
            continue

        m_sec = section_re.match(text)
        if m_sec:
            current_section_code = m_sec.group(1)
            current_section_title = m_sec.group(2).strip()
            if current_section_code and not current_chapter_code:
                current_chapter_code = current_section_code[:1]
                current_chapter_title = chapter_titles.get(current_chapter_code)
            current_question_number = None
            raw_rows.append({
                "page_num": idx + 1,
                "source_page_num": None,
                "source_para_num": idx + 1,
                "source_locator_type": "paragraph",
                "page_text": text,
                "content_role": "section_heading",
                "section_code": current_section_code,
                "section_title": current_section_title,
                "chapter_code": current_chapter_code,
                "chapter_title": current_chapter_title,
                "question_number": None,
            })
            continue

        m_q = question_re.match(text)
        if m_q and current_section_code:
            current_question_number = int(m_q.group(1))
            raw_rows.append({
                "page_num": idx + 1,
                "source_page_num": None,
                "source_para_num": idx + 1,
                "source_locator_type": "paragraph",
                "page_text": text,
                "content_role": "question",
                "section_code": current_section_code,
                "section_title": current_section_title,
                "chapter_code": current_chapter_code,
                "chapter_title": current_chapter_title,
                "question_number": current_question_number,
            })
            continue

        raw_rows.append({
            "page_num": idx + 1,
            "source_page_num": None,
            "source_para_num": idx + 1,
            "source_locator_type": "paragraph",
            "page_text": text,
            "content_role": "answer" if current_section_code else "narrative",
            "section_code": current_section_code,
            "section_title": current_section_title,
            "chapter_code": current_chapter_code,
            "chapter_title": current_chapter_title,
            "question_number": current_question_number,
        })

    section_text = {}
    for row in raw_rows:
        if row["section_code"] and row["content_role"] in {"answer", "narrative"}:
            section_text.setdefault(row["section_code"], []).append(row["page_text"])

    section_flags = {
        section_code: {
            "has_substantive_answer": has_substantive_answer(" ".join(texts)),
            "referenced_appendices": extract_appendix_references(" ".join(texts)),
        }
        for section_code, texts in section_text.items()
    }

    page_rows = []
    for row in raw_rows:
        flags = section_flags.get(row["section_code"], {"has_substantive_answer": False, "referenced_appendices": []})
        page_rows.append(_make_page_row(
            doc_meta=doc_meta,
            page_num=row["page_num"],
            page_text=row["page_text"],
            parse_method="python_docx_manager_ddq",
            content_role=row["content_role"],
            source_page_num=row.get("source_page_num"),
            source_para_num=row.get("source_para_num"),
            source_locator_type=row.get("source_locator_type") or "paragraph",
            section_code=row["section_code"],
            section_title=row["section_title"],
            chapter_code=row["chapter_code"],
            chapter_title=row["chapter_title"],
            question_number=row["question_number"],
            has_substantive_answer=bool(flags["has_substantive_answer"]),
            referenced_appendices=json.dumps(flags["referenced_appendices"], ensure_ascii=True),
        ))
    return page_rows, metadata


def _parse_generic_docx(local_path: str, doc_meta: dict) -> tuple[list[dict], dict]:
    doc = Document(local_path)
    page_rows = []
    for idx, para in enumerate(doc.paragraphs):
        text = normalize_text(para.text)
        if not text:
            continue
        page_rows.append(_make_page_row(
            doc_meta=doc_meta,
            page_num=idx + 1,
            page_text=text,
            parse_method="python_docx_generic",
            content_role="narrative",
            source_page_num=None,
            source_para_num=idx + 1,
            source_locator_type="paragraph",
        ))
    return page_rows, {}


def _parse_pdf(local_path: str, doc_meta: dict) -> tuple[list[dict], dict]:
    pdf = fitz.open(local_path)
    page_rows = []
    for idx in range(pdf.page_count):
        page = pdf.load_page(idx)
        text = page.get_text("text") or ""
        page_rows.append(_make_page_row(
            {
                **doc_meta,
                "file_path_local": local_path,
            },
            page_num=idx + 1,
            page_text=text,
            parse_method="pymupdf_text",
            content_role="page",
            source_page_num=idx + 1,
            source_para_num=idx + 1,
            source_locator_type="page",
        ))
    pdf.close()
    return page_rows, {}


page_rows = []
ok_doc_ids = []
failed_docs = []
skipped_template_ids = []
doc_metadata_updates = []

for d in docs:
    doc_id = d["document_id"]
    local_path = d["file_path_local"] or uc_dbfs_to_local_path(d["file_path_dbfs"])
    try:
        if d["source_role"] == "report_template":
            skipped_template_ids.append(doc_id)
            doc_metadata_updates.append({
                "document_id": doc_id,
                "manager_name": None,
                "mandate_name": None,
                "investment_strategy": None,
                "report_finalization_date": None,
                "portfolio_odd_managers": None,
                "authors": None,
            })
            continue

        if d["file_ext"] == ".pdf":
            parsed_rows, metadata = _parse_pdf(local_path, d)
            parse_method = "pymupdf_text"
        elif d["file_ext"] == ".docx" and d["source_role"] == "manager_completed_ddq":
            parsed_rows, metadata = _parse_manager_ddq(local_path, d)
            parse_method = "python_docx_manager_ddq"
        elif d["file_ext"] == ".docx":
            parsed_rows, metadata = _parse_generic_docx(local_path, d)
            parse_method = "python_docx_generic"
        else:
            raise ValueError(f"Unsupported file type: {d['file_name']}")

        page_rows.extend(parsed_rows)
        ok_doc_ids.append((doc_id, parse_method))
        doc_metadata_updates.append({"document_id": doc_id, **metadata})
    except Exception as e:
        failed_docs.append((doc_id, str(e)[:4000]))
        log_pipeline_error(
            ERRORS_TABLE,
            stage="parse_source",
            engagement_id=ENGAGEMENT_ID,
            document_id=doc_id,
            source_path=d["file_path_dbfs"],
            error=e,
            extra={"file_path_local": local_path, "source_role": d["source_role"]},
        )

print("Parsed OK docs:", len(ok_doc_ids))
print("Skipped template docs:", len(skipped_template_ids))
print("Failed docs:", len(failed_docs))

if not page_rows and not skipped_template_ids:
    raise RuntimeError("No source rows were parsed. Check file paths and source-role classification.")

# COMMAND ----------

pages_schema = T.StructType([
    T.StructField("engagement_id", T.StringType(), nullable=False),
    T.StructField("document_id", T.StringType(), nullable=False),
    T.StructField("file_name", T.StringType(), nullable=True),
    T.StructField("source_path_dbfs", T.StringType(), nullable=True),
    T.StructField("source_path_local", T.StringType(), nullable=True),
    T.StructField("page_num", T.IntegerType(), nullable=False),
    T.StructField("source_page_num", T.IntegerType(), nullable=True),
    T.StructField("source_para_num", T.IntegerType(), nullable=True),
    T.StructField("source_locator_type", T.StringType(), nullable=True),
    T.StructField("source_locator_label", T.StringType(), nullable=True),
    T.StructField("page_text", T.StringType(), nullable=True),
    T.StructField("page_char_count", T.IntegerType(), nullable=True),
    T.StructField("parse_method", T.StringType(), nullable=True),
    T.StructField("source_role", T.StringType(), nullable=True),
    T.StructField("source_tier", T.IntegerType(), nullable=True),
    T.StructField("section_code", T.StringType(), nullable=True),
    T.StructField("section_title", T.StringType(), nullable=True),
    T.StructField("chapter_code", T.StringType(), nullable=True),
    T.StructField("chapter_title", T.StringType(), nullable=True),
    T.StructField("content_role", T.StringType(), nullable=True),
    T.StructField("question_number", T.IntegerType(), nullable=True),
    T.StructField("has_substantive_answer", T.BooleanType(), nullable=True),
    T.StructField("referenced_appendices", T.StringType(), nullable=True),
])

if ok_doc_ids:
    spark.createDataFrame([(x[0],) for x in ok_doc_ids], ["document_id"]).createOrReplaceTempView("docs_to_replace_pages")
    spark.sql(f"""
    DELETE FROM {PAGES_TABLE}
    WHERE engagement_id = '{ENGAGEMENT_ID}'
      AND document_id IN (SELECT document_id FROM docs_to_replace_pages)
    """)

if page_rows:
    df_pages = spark.createDataFrame(page_rows, schema=pages_schema).withColumn("parse_ts", F.current_timestamp())
    df_pages.createOrReplaceTempView("new_pages")
    spark.sql(f"""
    MERGE INTO {PAGES_TABLE} t
    USING new_pages s
    ON t.engagement_id = s.engagement_id AND t.document_id = s.document_id AND t.page_num = s.page_num
    WHEN MATCHED THEN UPDATE SET *
    WHEN NOT MATCHED THEN INSERT *
    """)

if doc_metadata_updates:
    metadata_rows = []
    for item in doc_metadata_updates:
        metadata_rows.append({
            "document_id": item["document_id"],
            "manager_name": item.get("manager_name"),
            "mandate_name": item.get("mandate_name"),
            "investment_strategy": item.get("investment_strategy"),
            "report_finalization_date": item.get("report_finalization_date"),
            "portfolio_odd_managers": item.get("portfolio_odd_managers"),
            "authors": item.get("authors"),
        })
    metadata_schema = T.StructType([
        T.StructField("document_id", T.StringType(), nullable=False),
        T.StructField("manager_name", T.StringType(), nullable=True),
        T.StructField("mandate_name", T.StringType(), nullable=True),
        T.StructField("investment_strategy", T.StringType(), nullable=True),
        T.StructField("report_finalization_date", T.StringType(), nullable=True),
        T.StructField("portfolio_odd_managers", T.StringType(), nullable=True),
        T.StructField("authors", T.StringType(), nullable=True),
    ])
    spark.createDataFrame(metadata_rows, schema=metadata_schema).createOrReplaceTempView("doc_metadata_updates")
    spark.sql(f"""
    MERGE INTO {DOCUMENTS_TABLE} t
    USING doc_metadata_updates s
    ON t.document_id = s.document_id
    WHEN MATCHED THEN UPDATE SET
      t.manager_name = coalesce(s.manager_name, t.manager_name),
      t.mandate_name = coalesce(s.mandate_name, t.mandate_name),
      t.investment_strategy = coalesce(s.investment_strategy, t.investment_strategy),
      t.report_finalization_date = coalesce(s.report_finalization_date, t.report_finalization_date),
      t.portfolio_odd_managers = coalesce(s.portfolio_odd_managers, t.portfolio_odd_managers),
      t.authors = coalesce(s.authors, t.authors)
    """)

if ok_doc_ids:
    spark.createDataFrame(ok_doc_ids, ["document_id", "parse_method"]).createOrReplaceTempView("ok_docs")
    spark.sql(f"""
    MERGE INTO {DOCUMENTS_TABLE} t
    USING ok_docs s
    ON t.document_id = s.document_id
    WHEN MATCHED THEN UPDATE SET
      t.parse_status = 'done',
      t.parse_method = s.parse_method,
      t.parse_ts = current_timestamp(),
      t.parse_error = null
    """)

if skipped_template_ids:
    spark.createDataFrame([(x,) for x in skipped_template_ids], ["document_id"]).createOrReplaceTempView("skipped_templates")
    spark.sql(f"""
    MERGE INTO {DOCUMENTS_TABLE} t
    USING skipped_templates s
    ON t.document_id = s.document_id
    WHEN MATCHED THEN UPDATE SET
      t.parse_status = 'skipped',
      t.parse_method = 'report_template_deferred',
      t.parse_ts = current_timestamp(),
      t.parse_error = null,
      t.chunk_status = 'skipped',
      t.chunk_ts = current_timestamp(),
      t.chunk_error = null,
      t.index_status = 'skipped',
      t.index_ts = current_timestamp(),
      t.index_error = null
    """)

if failed_docs:
    spark.createDataFrame(failed_docs, ["document_id", "parse_error"]).createOrReplaceTempView("failed_docs")
    spark.sql(f"""
    MERGE INTO {DOCUMENTS_TABLE} t
    USING failed_docs s
    ON t.document_id = s.document_id
    WHEN MATCHED THEN UPDATE SET
      t.parse_status = 'error',
      t.parse_ts = current_timestamp(),
      t.parse_error = s.parse_error
    """)

# COMMAND ----------

display(
    spark.table(PAGES_TABLE)
    .filter(F.col("engagement_id") == ENGAGEMENT_ID)
    .groupBy("document_id", "file_name", "source_role")
    .agg(F.count("*").alias("rows"), F.sum("page_char_count").alias("chars"))
    .orderBy("source_role", "file_name")
)

show_validation_snapshot(CATALOG, SCHEMA, ENGAGEMENT_ID)
