# Databricks notebook source
# NOTEBOOK FILE: 06_parse_ddq.py
# COMMAND ----------
# MAGIC %pip install -U python-docx

# COMMAND ----------
dbutils.library.restartPython()

# COMMAND ----------
# MAGIC %run ./_config

# COMMAND ----------
# MAGIC %run ./_utils

# COMMAND ----------
# MAGIC %md
# MAGIC **Legacy Notebook**
# MAGIC
# MAGIC This notebook belongs to the deprecated DDQ-fill workflow. The default orchestrated path for
# MAGIC this repository is now the ODD report workflow (`06_parse_odd_report` -> `07_assess_odd_report`
# MAGIC -> `08_fill_odd_report`).
# MAGIC
# MAGIC ## 06 Parse the DDQ `.docx` into `ddq_questions`
# MAGIC
# MAGIC Walks the source docx using `python-docx`, extracts numbered questions and their
# MAGIC `<Provide your answer here.>` placeholder paragraphs, classifies each question as
# MAGIC yes_no / numeric / free_text, and writes the structured table.
# MAGIC
# MAGIC Tables are recorded with `question_type='skipped_table'` for audit but not answered (v1).

# COMMAND ----------

import re
from docx import Document

dbutils.widgets.text("engagement_id", "odd_ssga_2025")
dbutils.widgets.text("questionnaire_profile", "odd_ssga_v1")
dbutils.widgets.text("source_docx_name", "")  # optional override; otherwise auto-detect single .docx

ENGAGEMENT_ID = dbutils.widgets.get("engagement_id").strip()
PROFILE_KEY = dbutils.widgets.get("questionnaire_profile").strip()
SOURCE_DOCX_NAME = dbutils.widgets.get("source_docx_name").strip() or None

profile = QUESTIONNAIRE_PROFILES.get(PROFILE_KEY)
if profile is None:
    raise ValueError(f"Unknown questionnaire_profile '{PROFILE_KEY}'. Add it to _config.py.")

paths = engagement_paths(ENGAGEMENT_ID)
qdir_local = paths["questionnaire_local"]

# Locate the source docx.
import os
candidates = [f for f in os.listdir(qdir_local) if f.lower().endswith(".docx")]
print("Found docx candidates:", candidates)
if SOURCE_DOCX_NAME:
    if SOURCE_DOCX_NAME not in candidates:
        raise ValueError(f"source_docx_name='{SOURCE_DOCX_NAME}' not found in {qdir_local}")
    docx_name = SOURCE_DOCX_NAME
elif len(candidates) == 1:
    docx_name = candidates[0]
else:
    raise ValueError(
        f"Expected exactly 1 .docx in {qdir_local} (or set widget source_docx_name). Got: {candidates}"
    )

docx_local = os.path.join(qdir_local, docx_name)
docx_dbfs = f"{paths['questionnaire_dbfs']}/{docx_name}"
print("Parsing:", docx_local)

# COMMAND ----------

placeholder_re = re.compile(profile["answer_placeholder_pattern"])
section_re = re.compile(profile["section_heading_regex"])
qnum_re = re.compile(profile["question_numbering_regex"])

doc = Document(docx_local)

current_section_id = None
current_section_title = None
question_rows = []
pending_question = None  # (q_num, q_text, q_para_idx) waiting for next placeholder

for idx, para in enumerate(doc.paragraphs):
    text = (para.text or "").strip()
    if not text:
        continue

    # 1) Section heading?
    m_sec = section_re.match(text)
    if m_sec:
        current_section_id = m_sec.group(1)
        current_section_title = m_sec.group(2).strip()
        continue

    # 2) Numbered question?
    m_q = qnum_re.match(text)
    if m_q:
        q_num = int(m_q.group(1))
        q_text = m_q.group(2).strip()
        # If a previous question never found a placeholder, store it without one.
        if pending_question is not None:
            pn, pt, p_idx = pending_question
            question_rows.append({
                "section_id": current_section_id,
                "section_title": current_section_title,
                "question_number": pn,
                "question_text": pt,
                "placeholder_paragraph_index": None,
            })
        pending_question = (q_num, q_text, idx)
        continue

    # 3) Placeholder?
    if pending_question is not None and placeholder_re.search(text):
        q_num, q_text, _q_idx = pending_question
        question_rows.append({
            "section_id": current_section_id,
            "section_title": current_section_title,
            "question_number": q_num,
            "question_text": q_text,
            "placeholder_paragraph_index": idx,
        })
        pending_question = None

# Flush any trailing pending question without placeholder.
if pending_question is not None:
    pn, pt, _ = pending_question
    question_rows.append({
        "section_id": current_section_id,
        "section_title": current_section_title,
        "question_number": pn,
        "question_text": pt,
        "placeholder_paragraph_index": None,
    })

print(f"Extracted {len(question_rows)} questions ({sum(1 for r in question_rows if r['placeholder_paragraph_index'] is not None)} with placeholders).")

# Skipped-table audit rows (one row per table found).
skipped_table_rows = []
if profile.get("skip_tables", True):
    for ti, _tbl in enumerate(doc.tables):
        skipped_table_rows.append({
            "section_id": current_section_id,
            "section_title": current_section_title,
            "question_number": 10000 + ti,
            "question_text": f"<table {ti}>",
            "placeholder_paragraph_index": None,
        })
print(f"Skipped {len(skipped_table_rows)} tables (audit-only rows).")

# COMMAND ----------

now_ts = F.current_timestamp()

def _build_rows(rows: list[dict], qtype_override: str | None):
    out = []
    for r in rows:
        if qtype_override:
            qtype = qtype_override
        elif r["placeholder_paragraph_index"] is None:
            # In v1 we only fill explicit placeholder paragraphs. Questions without a placeholder
            # are treated as non-answerable structured items (for example embedded tables).
            qtype = "skipped_table"
        else:
            qtype = classify_question_type(r["question_text"], profile)
        qid = sha256_hex(f"{ENGAGEMENT_ID}||{PROFILE_KEY}||{r['question_number']}||{r['question_text']}")
        out.append({
            "engagement_id": ENGAGEMENT_ID,
            "question_id": qid,
            "questionnaire_profile": PROFILE_KEY,
            "section_id": r["section_id"],
            "section_title": r["section_title"],
            "question_number": int(r["question_number"]),
            "question_text": r["question_text"],
            "question_type": qtype,
            "placeholder_paragraph_index": r["placeholder_paragraph_index"],
            "source_docx_path": docx_dbfs,
        })
    return out


all_rows = _build_rows(question_rows, qtype_override=None) + _build_rows(skipped_table_rows, qtype_override="skipped_table")

if not all_rows:
    raise RuntimeError("No questions extracted; check the docx structure and profile regexes.")

q_schema = T.StructType([
    T.StructField("engagement_id", T.StringType()),
    T.StructField("question_id", T.StringType()),
    T.StructField("questionnaire_profile", T.StringType()),
    T.StructField("section_id", T.StringType()),
    T.StructField("section_title", T.StringType()),
    T.StructField("question_number", T.IntegerType()),
    T.StructField("question_text", T.StringType()),
    T.StructField("question_type", T.StringType()),
    T.StructField("placeholder_paragraph_index", T.IntegerType()),
    T.StructField("source_docx_path", T.StringType()),
])
df_q = (
    spark.createDataFrame(all_rows, schema=q_schema)
    .withColumn("load_ts", now_ts)
)
df_q.createOrReplaceTempView("new_questions")

# Replace this engagement's questions wholesale.
spark.sql(f"DELETE FROM {QUESTIONS_TABLE} WHERE engagement_id = '{ENGAGEMENT_ID}'")
spark.sql(f"INSERT INTO {QUESTIONS_TABLE} SELECT * FROM new_questions")

# COMMAND ----------

display(
    spark.table(QUESTIONS_TABLE)
    .filter(F.col("engagement_id") == ENGAGEMENT_ID)
    .groupBy("question_type")
    .count()
    .orderBy("question_type")
)

display(
    spark.table(QUESTIONS_TABLE)
    .filter(F.col("engagement_id") == ENGAGEMENT_ID)
    .filter(F.col("question_type") != "skipped_table")
    .orderBy("question_number")
    .limit(10)
)

show_validation_snapshot(CATALOG, SCHEMA, ENGAGEMENT_ID)
