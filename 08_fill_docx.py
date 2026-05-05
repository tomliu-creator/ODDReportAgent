# Databricks notebook source
# NOTEBOOK FILE: 08_fill_docx.py
# COMMAND ----------
# MAGIC %pip install -U python-docx

# COMMAND ----------
dbutils.library.restartPython()

# COMMAND ----------
# MAGIC %run ./_config

# COMMAND ----------
# MAGIC %run ./_utils

# COMMAND ----------
# MAGIC %md
# MAGIC ## 08 Fill the source `.docx` with draft answers + inline citations
# MAGIC
# MAGIC Walks `doc.paragraphs` and for each row in `ddq_questions` (excluding `skipped_table`)
# MAGIC replaces the `<Provide your answer here.>` placeholder at `placeholder_paragraph_index`
# MAGIC with the draft answer plus an italic-gray citation suffix:
# MAGIC `[Source: <file>, p.<N>; ...]`.
# MAGIC
# MAGIC Saves to `engagements/<engagement_id>/output/<original_basename>.filled.docx`.

# COMMAND ----------

import os
import shutil
import uuid
from docx import Document
from docx.shared import RGBColor

dbutils.widgets.text("engagement_id", "odd_ssga_2025")
dbutils.widgets.text("questionnaire_profile", "odd_ssga_v1")
ENGAGEMENT_ID = dbutils.widgets.get("engagement_id").strip()
PROFILE_NAME = dbutils.widgets.get("questionnaire_profile").strip() or "odd_ssga_v1"
PROFILE = get_questionnaire_profile(PROFILE_NAME)

paths = engagement_paths(ENGAGEMENT_ID)
qdir_local = paths["questionnaire_local"]
out_dir_local = paths["output_local"]
os.makedirs(out_dir_local, exist_ok=True)

# COMMAND ----------

# Find the source docx (single .docx in questionnaire/).
candidates = [f for f in os.listdir(qdir_local) if f.lower().endswith(".docx")]
if len(candidates) != 1:
    raise ValueError(f"Expected exactly 1 .docx in {qdir_local}; got: {candidates}")
docx_local = os.path.join(qdir_local, candidates[0])
out_basename = candidates[0].rsplit(".docx", 1)[0] + ".filled.docx"
out_local = os.path.join(out_dir_local, out_basename)
out_dbfs = f"{paths['output_dbfs']}/{out_basename}"
out_dbfs_local = out_dbfs.replace("dbfs:", "/dbfs")
workspace_output_dir = "/Workspace/Users/tomliushopping@gmail.com/notebooks_ddq/output_docs"
os.makedirs(workspace_output_dir, exist_ok=True)
out_workspace_local = os.path.join(workspace_output_dir, out_basename)
print("Source:", docx_local)
print("Output:", out_local)

# COMMAND ----------

q_rows = (
    spark.table(QUESTIONS_TABLE)
    .filter(F.col("engagement_id") == ENGAGEMENT_ID)
    .filter(F.col("question_type") != "skipped_table")
    .filter(F.col("placeholder_paragraph_index").isNotNull())
    .select("question_id", "question_number", "question_text", "placeholder_paragraph_index")
    .collect()
)
print("Answerable questions:", len(q_rows))

a_rows = (
    spark.table(ANSWERS_TABLE)
    .filter(F.col("engagement_id") == ENGAGEMENT_ID)
    .select("question_id", "draft_answer", "citations")
    .collect()
)
answers_by_id = {r["question_id"]: r.asDict() for r in a_rows}
missing = [r["question_id"] for r in q_rows if r["question_id"] not in answers_by_id]
if missing:
    print(f"WARNING: {len(missing)} questions have no answer (will leave placeholder).")


def _field(obj, name):
    if obj is None:
        return None
    if isinstance(obj, dict):
        return obj.get(name)
    try:
        return obj[name]
    except Exception:
        return getattr(obj, name, None)


all_cited_chunk_ids = sorted({
    _field(c, "chunk_id")
    for row in a_rows
    for c in (row["citations"] or [])
    if c and _field(c, "chunk_id")
})
chunk_meta_by_id = {}
if all_cited_chunk_ids:
    chunk_meta_rows = (
        spark.table(CHUNKS_TABLE)
        .filter(F.col("engagement_id") == ENGAGEMENT_ID)
        .filter(F.col("chunk_id").isin(all_cited_chunk_ids))
        .select("chunk_id", "file_name", "source_tier")
        .collect()
    )
    chunk_meta_by_id = {r["chunk_id"]: r.asDict() for r in chunk_meta_rows}

# COMMAND ----------

doc = Document(docx_local)


def _format_citations(cits) -> str:
    if not cits:
        return ""
    parts = []
    seen = set()
    for c in cits:
        f = _field(c, "file") or ""
        p = _field(c, "page")
        cid = _field(c, "chunk_id")
        meta = chunk_meta_by_id.get(cid) if cid else None
        display_name = get_doc_short_title((meta or {}).get("file_name") or f, PROFILE)
        tier = (meta or {}).get("source_tier")
        key = (display_name, p, tier)
        if key in seen:
            continue
        seen.add(key)
        if tier is not None:
            parts.append(f"{display_name} | p.{p} | tier {tier}")
        else:
            parts.append(f"{display_name} | p.{p}")
    if not parts:
        return ""
    return " [Source: " + "; ".join(parts) + "]"


filled = 0
for q in q_rows:
    p_idx = q["placeholder_paragraph_index"]
    if p_idx is None or p_idx >= len(doc.paragraphs):
        continue
    para = doc.paragraphs[p_idx]
    answer_row = answers_by_id.get(q["question_id"])
    if not answer_row:
        continue
    draft = (answer_row["draft_answer"] or "").strip()
    cit_suffix = _format_citations(answer_row["citations"])

    # Clear existing runs (placeholder text).
    for r in list(para.runs):
        r.text = ""
    # Replace with answer + citation suffix.
    answer_run = para.add_run(draft)
    if cit_suffix:
        cit_run = para.add_run(cit_suffix)
        cit_run.italic = True
        try:
            cit_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        except Exception:
            pass
    filled += 1

print(f"Filled {filled} placeholders.")

workspace_tmp_dir = "/Workspace/Users/tomliushopping@gmail.com/notebooks_ddq/.tmp"
os.makedirs(workspace_tmp_dir, exist_ok=True)
tmp_out_local = os.path.join(workspace_tmp_dir, f"{uuid.uuid4().hex}.docx")
doc.save(tmp_out_local)
final_output_path = out_dbfs
try:
    shutil.copyfile(tmp_out_local, out_dbfs_local)
except OSError as e:
    shutil.copyfile(tmp_out_local, out_workspace_local)
    final_output_path = out_workspace_local
    print(f"WARNING: could not write to volume output path ({out_dbfs_local}): {e}")
    print("Fell back to workspace-local output path.")
print("Saved:", final_output_path)

# COMMAND ----------

# Assertion: zero remaining placeholders.
import re
check_doc = Document(tmp_out_local)
remaining = sum(
    1 for p in check_doc.paragraphs
    if re.search(r"<Provide your answer here\.>", p.text or "")
)
print("Remaining `<Provide your answer here.>` placeholders:", remaining)
if remaining > 0:
    print("WARNING: some placeholders were not filled. Likely answers missing for those questions.")

print("Done. Output at:", final_output_path)
